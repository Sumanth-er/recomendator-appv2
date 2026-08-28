"""Word rendering of the approval package.

Same structure as the markdown version - both read render.package - so the
file a buyer forwards says exactly what the screen said. This is the format
the approval actually travels in: it gets signed, commented on and filed.

python-docx is imported inside the function rather than at module scope. The
rest of the app must still start and serve if the package is missing.
"""
from __future__ import annotations

import io

from .package import build_package

TITLE_SIZE = 20
SUBTITLE_SIZE = 13
BODY_SIZE = 10
NOTE_SIZE = 8.5


def _shade(cell, colour: str = "DCE6F1") -> None:
    """Header shading, which python-docx has no API for."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), colour)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_size(run, points: float, bold: bool = False, colour=None) -> None:
    from docx.shared import Pt

    run.font.size = Pt(points)
    run.bold = bold
    if colour is not None:
        run.font.color.rgb = colour


def _add_table(document, block: dict) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    header = block["header"]
    numeric = set(block.get("numeric") or [])

    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    for index, label in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(label))
        _set_size(run, BODY_SIZE, bold=True)
        _shade(cell)

    for row in block["rows"]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            run = paragraph.add_run(str(value))
            # The first column of a metric table is a label, not a number.
            _set_size(run, BODY_SIZE, bold=(index == 0 and len(header) and
                                            header[0] == ""))
            if index in numeric:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_paragraph()


def _add_fields(document, block: dict) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in block["rows"]:
        cells = table.add_row().cells
        cells[0].text = ""
        _set_size(cells[0].paragraphs[0].add_run(label), BODY_SIZE, bold=True)
        _shade(cells[0], "F2F5F9")
        cells[1].text = ""
        _set_size(cells[1].paragraphs[0].add_run(str(value or "")), BODY_SIZE)
    document.add_paragraph()


def _add_blocks(document, blocks: list[dict]) -> None:
    from docx.shared import RGBColor

    for block in blocks:
        kind = block["type"]
        if kind == "para":
            paragraph = document.add_paragraph()
            _set_size(paragraph.add_run(block["text"]), BODY_SIZE)
        elif kind == "note":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block["text"])
            _set_size(run, NOTE_SIZE, colour=RGBColor(0x60, 0x66, 0x70))
            run.italic = True
        elif kind == "bullets":
            for item in block["items"]:
                paragraph = document.add_paragraph(style="List Bullet")
                _set_size(paragraph.add_run(item), BODY_SIZE)
        elif kind == "table":
            _add_table(document, block)
        elif kind == "fields":
            _add_fields(document, block)
        elif kind == "signature":
            paragraph = document.add_paragraph()
            _set_size(paragraph.add_run(
                "Approver signature: ________________________"
                "          Date: ________________"), BODY_SIZE)


def _add_footer(section, text: str) -> None:
    """Footer with the document name and a real page-number field."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{text}    |    Page ")
    run.font.size = Pt(NOTE_SIZE)
    run.font.color.rgb = RGBColor(0x60, 0x66, 0x70)

    # PAGE is a field code; Word evaluates it when the document opens.
    field = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, end):
        field._r.append(element)
    field.font.size = Pt(NOTE_SIZE)
    field.font.color.rgb = RGBColor(0x60, 0x66, 0x70)


def render_docx(run_id: str) -> bytes | None:
    """The approval package as a .docx, or None when the run does not exist."""
    package = build_package(run_id)
    if package is None:
        return None

    from docx import Document
    from docx.shared import Pt, RGBColor

    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(BODY_SIZE)

    for section in document.sections:
        section.left_margin = section.right_margin = Pt(54)   # 0.75"
        _add_footer(section, package["footer"])

    title = document.add_paragraph()
    _set_size(title.add_run(package["title"]), TITLE_SIZE, bold=True)

    subtitle = document.add_paragraph()
    _set_size(subtitle.add_run(f"{package['subtitle']} — {package['date_text']}"),
              SUBTITLE_SIZE, bold=True, colour=RGBColor(0x1B, 0x4F, 0x8A))

    _add_fields(document, package["meta"])

    for section in package["sections"]:
        heading = document.add_paragraph()
        _set_size(heading.add_run(f"{section['number']}. {section['title']}"),
                  12, bold=True, colour=RGBColor(0x1B, 0x4F, 0x8A))
        _add_blocks(document, section["blocks"])

    trail = document.add_paragraph()
    run = trail.add_run(
        f"Generated by engine {package['engine_version']} on "
        f"{package['generated_at']}. Every figure above is either taken from a "
        f"supplier quotation or derived by the engine as labelled; no figure was "
        f"produced by a language model.")
    _set_size(run, NOTE_SIZE, colour=RGBColor(0x60, 0x66, 0x70))
    run.italic = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
